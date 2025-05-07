from docx import Document

# Create the document again
doc = Document()

# Title
doc.add_heading('Pre-Lab Questions', 0)

# Question 1
doc.add_heading('1. What is the moment of inertia of a homogeneous thin disc? Derive the formula.', level=1)
doc.add_paragraph(
    "For a homogeneous thin disc of mass M and radius R, rotating about an axis perpendicular to the disc and through its center, the moment of inertia I is given by:\n"
    "I = (1/2)MR²\n\n"
    "Derivation:\n"
    "The moment of inertia is defined as:\n"
    "I = ∫ r² dm\n\n"
    "where r is the distance from the axis of rotation, and dm is the infinitesimal mass element.\n"
    "Since the disc is uniform, the surface mass density σ is:\n"
    "σ = M / (πR²)\n\n"
    "Consider a small ring at radius r with thickness dr. The area of the ring is:\n"
    "dA = 2πr dr\n\n"
    "The mass of this ring is:\n"
    "dm = σ dA = σ (2πr dr)\n\n"
    "Substituting into the inertia integral:\n"
    "I = ∫₀ᴿ r² (σ 2πr dr) = 2πσ ∫₀ᴿ r³ dr\n"
    "I = 2πσ [r⁴/4]₀ᴿ = 2πσ (R⁴/4)\n\n"
    "Substituting σ = M / (πR²):\n"
    "I = 2π(M/πR²)(R⁴/4) = (MR²)/2\n\n"
    "Thus, the moment of inertia of a homogeneous thin disc is I = (1/2)MR²."
)

# Question 2
doc.add_heading('2. Explain why altering the axis of rotation of an object changes its moment of inertia.', level=1)
doc.add_paragraph(
    "The moment of inertia depends on the distribution of mass relative to the axis of rotation. "
    "When the axis is changed, the distances of mass elements from the new axis change.\n\n"
    "By definition:\n"
    "I = ∫ r² dm\n\n"
    "where r is the perpendicular distance from the mass element to the axis. "
    "If the axis moves farther from the mass, r increases and so does I. "
    "If the axis moves closer to the center of mass, I decreases.\n\n"
    "This is described by the Parallel Axis Theorem:\n"
    "I = I_cm + Md²\n\n"
    "where I_cm is the moment of inertia about the center of mass, and d is the distance between the center of mass and the new axis. "
    "Thus, altering the axis of rotation changes the distribution of mass relative to the axis, affecting the moment of inertia."
)

# Save the document
doc_path = '/mnt/data/Pre_Lab_Questions.docx'
doc.save(doc_path)

doc_path
